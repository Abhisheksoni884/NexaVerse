"""
Create test documents for NexaVerse ingestion testing
"""
from docx import Document
from docx.shared import Pt, Inches
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER

# Document 1: HR Policy (PDF)
def create_hr_policy():
    pdf_file = "HR_Policy_2026.pdf"
    doc = SimpleDocTemplate(pdf_file, pagesize=letter,
                           rightMargin=72, leftMargin=72,
                           topMargin=72, bottomMargin=18)
    
    Story = []
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='Justify', alignment=TA_JUSTIFY))
    
    # Title
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor='#2C5282',
        spaceAfter=30,
        alignment=TA_CENTER
    )
    
    Story.append(Paragraph("NexaVerse Corporation", title_style))
    Story.append(Paragraph("Human Resources Policy Manual", title_style))
    Story.append(Paragraph("Effective Date: January 1, 2026", styles['Normal']))
    Story.append(Spacer(1, 0.5*inch))
    
    # Content
    sections = [
        ("1. Leave Policy", 
         "Employees are entitled to 25 days of paid time off (PTO) per year, accrued at 2.08 days per month. "
         "PTO includes vacation days, personal days, and sick leave. Unused PTO can be carried over to the next "
         "year up to a maximum of 10 days. Additional sick leave is provided for serious medical conditions."),
        
        ("2. Remote Work Policy",
         "NexaVerse supports flexible work arrangements. Employees may work remotely up to 3 days per week after "
         "completing their probationary period. Fully remote positions are available for certain roles. All remote "
         "workers must maintain regular communication with their teams and be available during core business hours "
         "(10 AM - 3 PM local time)."),
        
        ("3. Professional Development",
         "The company allocates $2,500 per employee annually for professional development, including conferences, "
         "training courses, certifications, and educational materials. Employees are encouraged to pursue skills "
         "development aligned with their career goals and company needs."),
        
        ("4. Performance Reviews",
         "Performance evaluations are conducted quarterly. Employees receive feedback on their achievements, "
         "areas for improvement, and goal setting for the next period. Annual reviews determine salary adjustments "
         "and promotion eligibility."),
        
        ("5. Benefits Package",
         "NexaVerse offers comprehensive health insurance (medical, dental, vision), 401(k) retirement plan with "
         "5% company match, life insurance, disability coverage, and wellness programs including gym memberships "
         "and mental health support."),
    ]
    
    for title, content in sections:
        Story.append(Paragraph(title, styles['Heading2']))
        Story.append(Spacer(1, 0.2*inch))
        Story.append(Paragraph(content, styles['Justify']))
        Story.append(Spacer(1, 0.3*inch))
    
    doc.build(Story)
    print(f"✓ Created {pdf_file}")

# Document 2: Q1 Financial Report (PDF)
def create_financial_report():
    pdf_file = "Q1_2026_Financial_Report.pdf"
    doc = SimpleDocTemplate(pdf_file, pagesize=letter,
                           rightMargin=72, leftMargin=72,
                           topMargin=72, bottomMargin=18)
    
    Story = []
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=22,
        textColor='#2C5282',
        spaceAfter=20,
        alignment=TA_CENTER
    )
    
    Story.append(Paragraph("NexaVerse Corporation", title_style))
    Story.append(Paragraph("Q1 2026 Financial Report", title_style))
    Story.append(Spacer(1, 0.3*inch))
    
    sections = [
        ("Executive Summary",
         "NexaVerse achieved record revenue of $42.5M in Q1 2026, representing 28% year-over-year growth. "
         "Net profit margin improved to 18%, driven by operational efficiency and strong product demand. "
         "The company expanded its customer base by 350 new enterprise clients."),
        
        ("Revenue Breakdown",
         "Product revenue: $28.3M (67% of total), Services revenue: $10.2M (24%), Licensing: $4.0M (9%). "
         "Our SaaS products continue to lead growth, with 85% customer retention rate. The enterprise segment "
         "contributed 72% of total revenue, up from 65% in Q1 2025."),
        
        ("Operating Expenses",
         "Total operating expenses were $24.8M, including R&D ($9.2M), Sales & Marketing ($8.5M), and "
         "General & Administrative ($7.1M). R&D spending increased 15% to support our AI initiative and "
         "new product development pipeline."),
        
        ("Cash Position",
         "Cash and equivalents total $67.3M, providing strong liquidity for operations and strategic investments. "
         "The company maintains zero debt and generated $8.2M in operating cash flow during the quarter."),
        
        ("Outlook",
         "Management expects Q2 2026 revenue between $45M-$48M, with continued investment in product innovation "
         "and market expansion. The company plans to hire 50 additional employees across engineering and sales."),
    ]
    
    for title, content in sections:
        Story.append(Paragraph(title, styles['Heading2']))
        Story.append(Spacer(1, 0.15*inch))
        Story.append(Paragraph(content, styles['BodyText']))
        Story.append(Spacer(1, 0.25*inch))
    
    doc.build(Story)
    print(f"✓ Created {pdf_file}")

# Document 3: Product Roadmap (DOCX)
def create_product_roadmap():
    doc = Document()
    
    # Title
    title = doc.add_heading('NexaVerse Product Roadmap 2026', 0)
    title.alignment = 1  # Center
    
    doc.add_paragraph('Last Updated: July 2026')
    doc.add_paragraph()
    
    # Q3 2026
    doc.add_heading('Q3 2026 - AI-Powered Features', level=1)
    doc.add_paragraph(
        'Launch NexaVerse AI Assistant with natural language processing capabilities. '
        'This feature will enable users to query their document repositories using conversational AI, '
        'with support for multi-document analysis and automated summarization. Expected completion: September 2026.'
    )
    
    doc.add_heading('Enhanced Security & Compliance', level=2)
    doc.add_paragraph(
        'Implement SOC 2 Type II compliance, add advanced encryption for data at rest, '
        'and introduce granular role-based access control (RBAC) with custom permission policies. '
        'ISO 27001 certification process to begin.'
    )
    
    # Q4 2026
    doc.add_heading('Q4 2026 - Platform Expansion', level=1)
    doc.add_paragraph(
        'Mobile application launch for iOS and Android platforms. Native mobile experience with '
        'offline document access, push notifications, and biometric authentication. Beta testing starts October.'
    )
    
    doc.add_heading('API & Integrations', level=2)
    doc.add_paragraph(
        'Public API release with comprehensive documentation. Pre-built integrations with Salesforce, '
        'Microsoft Teams, Slack, and Google Workspace. Webhook support for real-time event notifications.'
    )
    
    # 2027
    doc.add_heading('2027 Initiatives', level=1)
    
    doc.add_heading('Advanced Analytics Dashboard', level=2)
    doc.add_paragraph(
        'Enterprise analytics platform with customizable dashboards, usage trends, cost optimization insights, '
        'and predictive analytics. Real-time collaboration metrics and team productivity scoring.'
    )
    
    doc.add_heading('Multi-Language Support', level=2)
    doc.add_paragraph(
        'Expand platform to support 15 languages including Spanish, French, German, Japanese, and Mandarin. '
        'AI models trained for multilingual document understanding and query processing.'
    )
    
    doc.add_heading('Workflow Automation', level=2)
    doc.add_paragraph(
        'Low-code workflow builder for document processing pipelines. Automated approval workflows, '
        'scheduled report generation, and trigger-based actions. Integration with popular automation tools.'
    )
    
    doc.save('test_documents/Product_Roadmap_2026.docx')
    print("✓ Created Product_Roadmap_2026.docx")

# Document 4: Technical Architecture (PDF)
def create_technical_doc():
    pdf_file = "Technical_Architecture_Guide.pdf"
    doc = SimpleDocTemplate(pdf_file, pagesize=letter,
                           rightMargin=72, leftMargin=72,
                           topMargin=72, bottomMargin=18)
    
    Story = []
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=20,
        textColor='#2C5282',
        spaceAfter=20,
        alignment=TA_CENTER
    )
    
    Story.append(Paragraph("NexaVerse Technical Architecture", title_style))
    Story.append(Paragraph("System Design & Infrastructure Guide", styles['Heading3']))
    Story.append(Spacer(1, 0.3*inch))
    
    sections = [
        ("Architecture Overview",
         "NexaVerse is built on a modern microservices architecture deployed on Microsoft Azure. "
         "The platform leverages Azure OpenAI for AI capabilities, Azure AI Search for hybrid vector search, "
         "Cosmos DB for document storage, and Azure Blob Storage for file management. The frontend is built "
         "with React and TypeScript, while the backend uses Python FastAPI."),
        
        ("Document Processing Pipeline",
         "When a document is uploaded: (1) File stored in Azure Blob Storage, (2) Azure Document Intelligence "
         "extracts text and structure, (3) Content is chunked using token-aware algorithms with 500 token chunks "
         "and 50 token overlap, (4) Azure OpenAI generates embeddings using text-embedding-3-small model, "
         "(5) Chunks indexed in Azure AI Search with vector and keyword search capabilities."),
        
        ("RAG Query Flow",
         "User queries follow this path: (1) Query received via FastAPI endpoint, (2) Content Safety API validates "
         "input, (3) Query embedding generated, (4) Hybrid search against Azure AI Search (vector + keyword), "
         "(5) Top K results retrieved with relevance scores, (6) Context sent to GPT-4o with system prompt, "
         "(7) Response streamed to client via Server-Sent Events (SSE), (8) Citations included with chunk metadata."),
        
        ("Security & Access Control",
         "Authentication uses JWT tokens with HS256 signing. Role-based access control (RBAC) enforces permissions "
         "at three levels: API endpoints, search index filtering, and document visibility. All API calls are logged "
         "to Cosmos DB audit container. Content Safety API scans both inputs and outputs for policy violations."),
        
        ("Scalability & Performance",
         "The system auto-scales based on demand using Azure App Service. Azure AI Search handles 100+ concurrent "
         "queries with <200ms latency. Document processing is asynchronous with status polling. Response caching "
         "reduces OpenAI API costs. CDN serves static assets globally for optimal performance."),
    ]
    
    for title, content in sections:
        Story.append(Paragraph(title, styles['Heading2']))
        Story.append(Spacer(1, 0.15*inch))
        Story.append(Paragraph(content, styles['BodyText']))
        Story.append(Spacer(1, 0.25*inch))
    
    doc.build(Story)
    print(f"✓ Created {pdf_file}")

# Document 5: Sales Playbook (DOCX)
def create_sales_playbook():
    doc = Document()
    
    title = doc.add_heading('NexaVerse Sales Playbook', 0)
    title.alignment = 1
    
    doc.add_paragraph('Q3 2026 Edition')
    doc.add_paragraph()
    
    doc.add_heading('Target Customer Profile', level=1)
    doc.add_paragraph(
        'Ideal customers are mid-to-large enterprises (500+ employees) with significant document management needs. '
        'Key verticals include: Financial Services, Healthcare, Legal, Manufacturing, and Technology. '
        'Decision-makers typically include CTO, CIO, VP of Operations, and Compliance Officers.'
    )
    
    doc.add_heading('Value Proposition', level=1)
    doc.add_paragraph(
        'NexaVerse reduces document retrieval time by 80% through AI-powered search. '
        'ROI is typically realized within 6 months through increased employee productivity and reduced '
        'manual search efforts. Key differentiators: enterprise-grade security, hybrid search technology, '
        'and seamless Azure integration for existing Azure customers.'
    )
    
    doc.add_heading('Pricing Model', level=1)
    doc.add_paragraph('Starter Plan: $2,500/month - Up to 50 users, 10,000 documents, 5GB storage')
    doc.add_paragraph('Professional Plan: $7,500/month - Up to 250 users, 50,000 documents, 50GB storage')
    doc.add_paragraph('Enterprise Plan: Custom pricing - Unlimited users and documents, dedicated support, SLA guarantees')
    doc.add_paragraph()
    doc.add_paragraph('All plans include: AI-powered search, role-based access control, audit logging, and API access')
    
    doc.add_heading('Common Objections & Responses', level=1)
    
    doc.add_heading('Objection: "We already use SharePoint/Google Drive"', level=2)
    doc.add_paragraph(
        'Response: NexaVerse complements existing storage solutions by adding AI-powered search capabilities. '
        'While SharePoint stores files, NexaVerse makes content instantly discoverable through natural language queries. '
        'We integrate with SharePoint, so your team continues using familiar tools.'
    )
    
    doc.add_heading('Objection: "What about data security?"', level=2)
    doc.add_paragraph(
        'Response: Security is our top priority. All data is encrypted in transit (TLS 1.3) and at rest (AES-256). '
        'We are SOC 2 Type II certified and pursuing ISO 27001. Data resides in your chosen Azure region and never '
        'leaves your geography. We support SSO, MFA, and granular RBAC.'
    )
    
    doc.add_heading('Success Stories', level=1)
    doc.add_paragraph(
        'TechCorp Inc. (5,000 employees): Reduced support ticket resolution time by 60% by enabling support staff '
        'to instantly find product documentation and troubleshooting guides. ROI achieved in 4 months.'
    )
    doc.add_paragraph()
    doc.add_paragraph(
        'FinServe Global (2,000 employees): Compliance team reduced audit preparation time from 6 weeks to 3 days '
        'by quickly retrieving historical policies and regulatory documents. Saved $400K annually in compliance costs.'
    )
    
    doc.save('test_documents/Sales_Playbook_Q3_2026.docx')
    print("✓ Created Sales_Playbook_Q3_2026.docx")

if __name__ == "__main__":
    print("Creating test documents for NexaVerse...")
    print()
    
    try:
        create_hr_policy()
        create_financial_report()
        create_product_roadmap()
        create_technical_doc()
        create_sales_playbook()
        print()
        print("✅ All test documents created successfully!")
        print("📁 Location: test_documents/")
    except Exception as e:
        print(f"❌ Error creating documents: {e}")
        import traceback
        traceback.print_exc()
